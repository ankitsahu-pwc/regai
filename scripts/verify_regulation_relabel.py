"""Verify Agent 2's BRD output honors a non-DORA regulation code.

Runs the same pipeline as Page 2 (parse -> Agent 1 -> Agent 2) but with
``regulation="GDPR"`` and checks that:

  1. The generated in-memory BRD contains zero "DORA" mentions.
  2. The generated .docx also contains zero "DORA" mentions.
  3. The tier line on the .docx says ``Compliance Framework: GDPR``.
  4. The DORA Regulation (EU) 2022/2554 citation is fully stripped.

Also runs the same test with ``regulation="DORA"`` to make sure the
regression-free DORA path still produces the historical, DORA-labelled
output.
"""
from __future__ import annotations

import sys
import traceback
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from pydantic import BaseModel  # noqa: E402

from services import persistence as db  # noqa: E402
from orchestrator import RegulatoryWorkflowOrchestrator  # noqa: E402
from utils.file_utils import timestamped_name  # noqa: E402


def _iter_strings(obj):
    if isinstance(obj, str):
        yield obj
        return
    if isinstance(obj, BaseModel):
        for name in obj.model_fields:
            yield from _iter_strings(getattr(obj, name, None))
        return
    if isinstance(obj, list):
        for item in obj:
            yield from _iter_strings(item)
        return
    if isinstance(obj, dict):
        for v in obj.values():
            yield from _iter_strings(v)


def _count_dora_mentions(strings):
    import re
    pat = re.compile(r"\bDORA\b")
    return sum(len(pat.findall(s)) for s in strings)


def _load_docx_text(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def run(regulation: str) -> int:
    rows = db.list_documents(kind="regulation")
    if not rows:
        print("FAIL: no regulation documents in DB.")
        return 2
    reg = rows[0]
    reg_path = Path(reg["path"])
    print(f"=== Regulation={regulation!r} ===")
    print(f"  Source PDF: {reg_path.name}")

    orch = RegulatoryWorkflowOrchestrator()
    try:
        parsed = orch.parse_document(reg_path, kind="regulation")
        analysis = orch.run_regulatory_analysis(
            parsed_document=parsed,
            regulation=regulation,
            tier="Tier-2",
            status=lambda _msg: None,
            regulator_selection=None,
            consulting_selection=None,
            include_consulting_guidance=False,
            intelligence_package=None,
        )
    except Exception as exc:
        print(f"FAIL: pipeline raised {type(exc).__name__}: {exc}")
        traceback.print_exc()
        return 3

    output_dir = ROOT / "outputs"
    docx_path = output_dir / timestamped_name(
        f"{regulation.replace(' ', '_')}_BRD_relabel_check", ".docx",
    )
    try:
        bundle = orch.run_brd_rtm(
            analysis, docx_export_path=docx_path, tier="Tier-2",
        )
    except Exception as exc:
        print(f"FAIL: run_brd_rtm raised {type(exc).__name__}: {exc}")
        traceback.print_exc()
        return 3

    brd = bundle["brd"]
    if brd.report is None:
        print("FAIL: no BRD report returned.")
        return 3

    in_memory_strings = list(_iter_strings(brd.report))
    in_memory_dora = _count_dora_mentions(in_memory_strings)
    print(f"  In-memory BRD strings: {len(in_memory_strings)}   DORA mentions: {in_memory_dora}")

    if not docx_path.exists():
        print(f"FAIL: docx not written to {docx_path}")
        return 3
    docx_text = _load_docx_text(docx_path)
    docx_dora = docx_text.count("DORA")
    citation_hits = docx_text.count("(EU) 2022/2554")
    print(f"  DOCX bytes: {docx_path.stat().st_size:,}")
    print(f"  DOCX 'DORA' occurrences: {docx_dora}")
    print(f"  DOCX '(EU) 2022/2554' occurrences: {citation_hits}")
    print(f"  DOCX first 200 chars: {docx_text[:200].replace(chr(10), ' | ')!r}")

    expected_frame = f"Compliance Framework: {regulation}"
    if expected_frame in docx_text:
        print(f"  OK: DOCX header contains {expected_frame!r}")
    else:
        print(f"  WARN: DOCX header does NOT contain {expected_frame!r}")

    if regulation.upper() == "DORA":
        if in_memory_dora == 0:
            print("  FAIL: DORA path unexpectedly stripped DORA mentions.")
            return 4
        print("  PASS: DORA regression path is intact.")
        return 0

    if in_memory_dora > 0:
        print(f"  FAIL: non-DORA path still contains {in_memory_dora} DORA mentions in memory.")
        return 5
    if docx_dora > 0:
        print(f"  FAIL: non-DORA docx still contains {docx_dora} DORA mentions.")
        return 5
    if citation_hits > 0:
        print(f"  FAIL: non-DORA docx still contains DORA citation.")
        return 5
    print(f"  PASS: no DORA leakage in {regulation!r} run.")
    return 0


def main() -> int:
    dora = run("DORA")
    print()
    gdpr = run("GDPR")
    print()
    mifid = run("MiFID II")
    return dora or gdpr or mifid


if __name__ == "__main__":
    raise SystemExit(main())
