"""Generic DOCX parsing helpers.

This module is intentionally domain-agnostic. It extracts paragraphs and tables
in document order and exposes small text-normalisation helpers. BRD/FRD
specific logic (section prefixing, theme inference, confidence clamping) lives
in ``services/questionnaire_generator.py`` in Phase 5 and consumes the output
of these helpers.

The body-iteration trick (yielding paragraphs and tables in true document
order) was lifted from the original ``read_docx_requirements`` in
``generate_brd_questionnaire_streamlit_v11.py`` and generalised here.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Iterator, List, Tuple, Union

from docx import Document
from docx.document import Document as _DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

DocxSource = Union[str, Path, bytes, io.IOBase]


def clean_text(value: object) -> str:
    """Collapse whitespace and strip non-breaking spaces from any cell/paragraph value."""
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).replace("\u00a0", " ")).strip()


def normalise_header(value: str) -> str:
    """Lower-case + alphanumeric-only header key, e.g. ``DORA Alignment`` -> ``doraalignment``."""
    return re.sub(r"[^a-z0-9]+", "", (value or "").lower())


def _open_document(source: DocxSource) -> _DocxDocument:
    """Open a python-docx ``Document`` from a path, raw bytes, or file-like object."""
    if isinstance(source, (str, Path)):
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")
        return Document(str(path))
    if isinstance(source, bytes):
        return Document(io.BytesIO(source))
    if hasattr(source, "read"):
        return Document(source)
    raise TypeError(f"Unsupported DOCX source type: {type(source).__name__}")


def iter_body_blocks(doc: _DocxDocument) -> Iterator[Union[Paragraph, Table]]:
    """Yield paragraphs and tables in true document order.

    python-docx exposes ``doc.paragraphs`` and ``doc.tables`` separately, which
    breaks the parent/child relationship between section headings and the
    tables that follow them. Iterating the body's XML children preserves order.
    """
    paragraphs_by_element = {p._p: p for p in doc.paragraphs}
    tables_by_element = {t._tbl: t for t in doc.tables}
    for child in doc.element.body.iterchildren():
        if child in paragraphs_by_element:
            yield paragraphs_by_element[child]
        elif child in tables_by_element:
            yield tables_by_element[child]


def extract_paragraphs(source: DocxSource) -> List[str]:
    """Return a list of non-empty paragraph texts, in document order."""
    doc = _open_document(source)
    return [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]


def extract_tables(source: DocxSource) -> List[List[List[str]]]:
    """Return all tables as nested lists ``tables[table_idx][row_idx][col_idx]``."""
    doc = _open_document(source)
    tables: List[List[List[str]]] = []
    for table in doc.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            rows.append([clean_text(cell.text) for cell in row.cells])
        tables.append(rows)
    return tables


def iter_sectioned_tables(
    source: DocxSource,
) -> Iterator[Tuple[str, List[List[str]]]]:
    """Yield ``(preceding_section_heading, table_rows)`` pairs in document order.

    A "section heading" is detected as a paragraph beginning with a numeric
    label such as ``1.``, ``7.1.``, or ``10.2.3``. The latest seen heading is
    associated with every subsequent table until another heading appears.

    This is the building block the BRD requirement extractor uses to know
    which 7.x section a requirement table belongs to.
    """
    doc = _open_document(source)
    current_section = "Unspecified"
    section_pattern = re.compile(r"^\d+(\.\d+)*\.?\s+")
    for block in iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            text = clean_text(block.text)
            if text and section_pattern.match(text):
                current_section = text
        elif isinstance(block, Table):
            rows = [[clean_text(cell.text) for cell in row.cells] for row in block.rows]
            if rows:
                yield current_section, rows


def extract_full_text(source: DocxSource, include_tables: bool = True) -> str:
    """Return a single string containing paragraphs and (optionally) flattened table cells."""
    doc = _open_document(source)
    chunks: List[str] = []
    for block in iter_body_blocks(doc):
        if isinstance(block, Paragraph):
            text = clean_text(block.text)
            if text:
                chunks.append(text)
        elif isinstance(block, Table) and include_tables:
            for row in block.rows:
                cells = [clean_text(cell.text) for cell in row.cells]
                joined = " | ".join(c for c in cells if c)
                if joined:
                    chunks.append(joined)
    return "\n".join(chunks)


__all__ = [
    "DocxSource",
    "clean_text",
    "extract_full_text",
    "extract_paragraphs",
    "extract_tables",
    "iter_body_blocks",
    "iter_sectioned_tables",
    "normalise_header",
]
